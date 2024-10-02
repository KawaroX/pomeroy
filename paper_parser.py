import re
import argparse
from docx import Document

class PaperParser:
    def __init__(self):
        pass

    @staticmethod
    def safe_extract(pattern, text, default="未找到", flags=0):
        match = re.search(pattern, text, flags)
        return match.group(1).strip() if match else default

    @staticmethod
    def extract_paper_info(text):
        result = {}

        # 提取标题（第一行）
        title = text.split('\n')[0].strip()
        result["标题"] = title  # .append(("标题", title))

        # 提取作者和来源
        author = PaperParser.safe_extract(r'【作\s*者】\s*(.+)', text, "未找到作者")
        result["作者"] = author  # .append(("作者", author))

        source = PaperParser.safe_extract(r'【来\s*源】\s*(.+)', text, "未找到来源")
        result["来源"] = source  # ..append(("来源", source))

        # 提取目次
        toc_match = re.search(r'目\s*次([\s\S]+?)(?=摘\s*要|$)', text)
        if toc_match:
            toc_content = toc_match.group(1).strip()
            toc_items = re.findall(r'^([一二三四五六七八九十]+、.+)', toc_content, re.MULTILINE)
            max_len = max(len(item) for item in toc_items) if toc_items else 0
            result["目次"] = {"m_l": max_len,
                              "items": toc_items}  # .append(("目次", {"m_l": max_len, "items": toc_items}))

        # 提取摘要
        abstract = PaperParser.safe_extract(r'摘\s*要：([\s\S]+?)(?=关键词|$)', text, "未找到摘要")
        result["摘要"] = abstract  # .append(("摘要", abstract))

        # 提取关键词
        keywords = PaperParser.safe_extract(r'关键词：(.+)', text, "未找到关键词")
        result["关键词"] = keywords  # .append(("关键词", keywords))

        # 提取正文结构
        keyword_pos = text.find('关键词：')
        keyword_pos_end = text.find('目  次')
        if keyword_pos == -1:
            result["正文"] = [("e", "无法提取正文")]  # .append(("正文", "无法提取正文"))
        else:
            cont3nt = text[keyword_pos:keyword_pos_end].split('\n', 1)[1].strip()
            result["正文"] = PaperParser.extract_content_structure(cont3nt)  # .extend(extract_content_structure(content))
        print(result["正文"])
        # # 根据正文中的一级标题获取目次
        # result["目次"] = {"m_l": 0, "items": []}
        # for item in result["正文"]:
        #     if item[0] == "1":
        #         result["目次"]["items"].append(item[1])
        #         result["目次"]["m_l"] = max(result["目次"]["m_l"], len(item[1]))


        return result

    @staticmethod
    def extract_content_structure(content):
        structure = []
        lines = content.split('\n')
        current_text = []
        # current_level = 0

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if re.match(r'^[一二三四五六七八九十]+、', line):
                if current_text:
                    structure.append(("0", '\n'.join(current_text)))
                    current_text = []
                structure.append(("1", line))
                # current_level = 1
            elif re.match(r'^（[一二三四五六七八九十]+）', line):
                if current_text:
                    structure.append(("0", '\n'.join(current_text)))
                    current_text = []
                structure.append(("2", line))
                # current_level = 2
            elif re.match(r'^\d+\.', line):
                if current_text:
                    structure.append(("0", '\n'.join(current_text)))
                    current_text = []
                structure.append(("3", line))
                # current_level = 3
            else:
                current_text.append(line)

        if current_text:
            structure.append(("0", '\n'.join(current_text)))

        return structure

    @staticmethod
    def get_placeholder(diff, base="北方法学"):
        return ''.join((base * (diff // len(base) + 1))[:diff])

    @staticmethod
    def generate_html(paper_info):

        # 目次
        placeholder_base = "北方法学"
        for i in range(len(paper_info["目次"]["items"])):
            placeholder_index = 0
            title = paper_info["目次"]["items"][i]

            diff = paper_info["目次"]["m_l"] - len(title)
            placeholder = PaperParser.get_placeholder(diff, placeholder_base[placeholder_index:] + placeholder_base[:placeholder_index])

            paper_info["目次"]["items"][i] = f"""
                <section style="margin-top: 5px;">
                    <section style="text-align: center;font-size: 15px;line-height: 1.8;">
                        <p>{title}<span style="color: rgba(255, 255, 255, 0);">{placeholder}</span></p>
                    </section>
                </section>
            """
        print(paper_info["目次"])

        # 正文
        content_new = ""
        paragraph = []
        for _, (i, content) in enumerate(paper_info["正文"]):
            if content == "无法提取正文":
                print(f"无法提取正文 {i}")
            elif i == "1":
                content = content.split("、")
                content_new += """\n<section style=""><section style="font-size: 14px;color: rgb(62, 62, 62);letter-spacing: 1px;line-height: 2;padding-right: 6px;padding-left: 6px;">"""
                content_new += "\n".join(paragraph) if paragraph else ""
                content_new += "\n</section></section>"
                content_new += f"""
                <section style="margin-top: 10px;margin-bottom: 10px;">
                    <section style="display: flex;width: 100%;flex-flow: column;">
                        <section style="z-index: 1;">
                            <section style="text-align: left;justify-content: flex-start;display: flex;flex-flow: row;">
                                <section style="display: inline-block;vertical-align: middle;width: auto;align-self: center;flex: 0 0 auto;min-width: 5%;height: auto;margin-right: -18px;">
                                    <section style="display: flex;width: 100%;flex-flow: column;">
                                        <section style="z-index: 1;">
                                            <section style="justify-content: flex-start;display: flex;flex-flow: row;">
                                                <section style="display: inline-block;vertical-align: middle;width: auto;min-width: 5%;flex: 0 0 auto;height: auto;align-self: center;">
                                                    <section style="font-size: 19px;margin-right: 0%;margin-left: 0%;text-align: center;">
                                                        <section style="display: inline-block;border-width: 1px;border-style: solid;border-color: rgb(190, 167, 139);background-color: rgb(190, 167, 139);width: 1.8em;height: 1.8em;line-height: 1.8em;border-radius: 100%;margin-left: auto;margin-right: auto;font-size: 18px;color: rgb(255, 255, 255);"><p><strong>{content[0]}</strong></p>
                                                        </section>
                                                    </section>
                                                </section>
                                                <section style="display: inline-block;vertical-align: middle;width: auto;align-self: center;flex: 0 0 auto;min-width: 5%;height: auto;line-height: 0.1;">
                                                    <section style="transform: translate3d(-1px, 0px, 0px);">
                                                        <section style="display: inline-block;width: 0px;height: 0px;vertical-align: top;overflow: hidden;border-style: solid;border-width: 4px 0px 4px 7px;border-color: rgba(255, 255, 255, 0) rgb(62, 62, 62) rgba(255, 255, 255, 0) rgb(190, 167, 139);">
                                                            <section style="text-align: justify;">
                                                                <p style="text-wrap: wrap;"><br></p>
                                                            </section>
                                                        </section>
                                                    </section>
                                                </section>
                                            </section>
                                        </section>
                                    </section>
                                </section>
                                <section style="display: inline-block;vertical-align: middle;width: auto;align-self: center;flex: 100 100 0%;height: auto;padding: 9px 10px 9px 33px;background-color: rgb(232, 228, 216);">
                                    <section style="text-align: justify;font-size: 18px;color: rgb(0, 0, 0);">
                                        <p style="text-wrap: wrap;">{content[1]}</p>
                                    </section>
                                </section>
                            </section>
                        </section>
                    </section>
                </section>
                """
                # print("!!!", paragraph, content_new)
                paragraph = []
                # print("end")
            else:
                if i == "2":
                    p = f"""<p style="text-indent: 2.1429em;margin-bottom: 9px;text-wrap: wrap;"><strong><span style="font-size: 16px;">{content}</span></strong></p>"""
                    paragraph.append(p)
                elif i == "3":
                    p = f"""<p style="text-indent: 2.1429em;margin-bottom: 9px;text-wrap: wrap;"><strong>{content}</strong></p>"""
                    paragraph.append(p)
                else:
                    content = content.split("\n")
                    for j in content:
                        if "* 因原文篇幅较长，已略去注释、图表等。" in j:
                            paragraph.append(
                                """<p style="text-indent: 2.1429em;margin-bottom: 9px;text-wrap: wrap;"><br></p><p style="text-wrap: wrap;"><strong>* 因原文篇幅较长，已略去注释、图表等。</strong></p>""")
                            break
                        else:
                            p = f"""<p style="text-indent: 2em;margin-bottom: 9px;text-wrap: wrap;">{j}</p>"""
                            paragraph.append(p)
            # print(i, content)
        content_new += """\n<section style=""><section style="font-size: 14px;color: rgb(62, 62, 62);letter-spacing: 1px;line-height: 2;padding-right: 6px;padding-left: 6px;">"""
        content_new += "\n".join(paragraph) if paragraph else ""
        content_new += "\n</section></section>"

        paper_info["正文"] = content_new

        output_body = f"""
        <body class="view rich_media_content autoTypeSetting24psection"
              lang="en"
              contenteditable="true"
              spellcheck="false"
              style="overflow-y: hidden; height: 15078px; cursor: text;">
            <section style="font-size: 16px;">
                <p style="text-align: center;">
                    <img class="rich_pages wxw-img"
                         data-galleryid=""
                         data-imgfileid="100000294"
                         data-ratio="0.28125"
                         data-type="gif"
                         data-w="960"
                         style="height: auto !important;"
                         src="https://mmbiz.qpic.cn/mmbiz_gif/iab7cNlld8FfILpZ5bZpvk6bwUtn2K8EdibwTU0d4ofpraMKRyhibgBF3OhmguCKibTwhdqFNMWfjAZRkQtRefKyzQ/640?wx_fmt=gif&amp;from=appmsg"
                         data-imgqrcoded="1">
                </p>
                <section style="margin-top: -5px;margin-bottom: 10px;">
                    <section style="font-size: 14px;color: rgb(62, 62, 62);line-height: 1.8;padding-right: 9px;padding-left: 9px;">
                        <p style="margin-bottom: 8px;text-indent: 0em;text-wrap: wrap;"><span style="text-indent: 2em;"><br></span></p>
                        <p style="margin-bottom: 8px;text-indent: 0em;text-wrap: wrap;"><span style="text-indent: 2em;">【作&nbsp;者】</span>{paper_info["作者"]}</p>
                        <p style="text-indent: 0em;text-wrap: wrap;">【来&nbsp;源】{paper_info["来源"]}</p>
                    </section>
                </section>
                <p style="text-wrap: wrap;"><br></p>
                <section style="text-align: left; justify-content: flex-start; display: flex; flex-flow: row; position: static; box-sizing: border-box;"><section style="display: inline-block; vertical-align: top; width: 24%; align-self: flex-start; flex: 0 0 auto; box-sizing: border-box;"><section style="text-align: center; margin-top: 10px; margin-bottom: 10px; line-height: 0; position: static; box-sizing: border-box;"><section style="max-width: 100%; vertical-align: middle; display: inline-block; line-height: 0; box-sizing: border-box;"><img class="raw-image" style="vertical-align: middle; max-width: 100%; width: 100%; box-sizing: border-box;" data-s="300,640" crossorigin="anonymous" src="https://img.xiumi.us/xmi/ua/3Mvma/i/28ef4f9b5fc04c955bb1cf23cd1e01db-sz_1737.gif"></section></section></section><section style="display: inline-block; vertical-align: middle; width: 52%; align-self: center; flex: 0 0 auto; background-position: 50% 50%; background-repeat: no-repeat; background-attachment: scroll; border-width: 0px; border-radius: 15px; border-style: none; border-color: rgb(62, 62, 62); overflow: hidden; background-image: url(&quot;https://img.xiumi.us/xmi/ua/3Mvma/i/2ce5101d7602568ce3a052e19e512ad6-sz_2126.gif&quot;); background-size: cover !important; box-sizing: border-box;"><section style="text-align: center; letter-spacing: 2px; box-sizing: border-box;"><p style="margin: 0px; padding: 0px; box-sizing: border-box;"><b style="box-sizing: border-box;">目次</b></p></section></section><section style="display: inline-block; vertical-align: top; width: 24%; align-self: flex-start; flex: 0 0 auto; box-sizing: border-box;"><section style="text-align: center; margin-top: 10px; margin-bottom: 10px; line-height: 0; position: static; box-sizing: border-box;"><section style="max-width: 100%; vertical-align: middle; display: inline-block; line-height: 0; box-sizing: border-box;"><img class="raw-image" style="vertical-align: middle; max-width: 100%; width: 100%; box-sizing: border-box;" data-s="300,640" crossorigin="anonymous" src="https://img.xiumi.us/xmi/ua/3Mvma/i/fb73f84eb3a373c4860cb54799e566a5-sz_1737.gif"></section></section></section></section>
                <p style="text-wrap: wrap;"><br></p>{"".join(paper_info["目次"]["items"])}<section style="text-align: left;justify-content: flex-start;display: flex;flex-flow: row;margin-top: 10px;margin-bottom: 20px;">
                    <section style="display: inline-block;width: auto;vertical-align: top;align-self: flex-start;flex: 0 0 auto;border-style: solid;border-width: 1px;min-width: 5%;height: auto;box-shadow: rgb(232, 228, 216) 6px 6px 0px 0px;padding: 23px;">
                        <section style="text-align: justify;font-size: 14px;color: rgb(62, 62, 62);letter-spacing: 1px;line-height: 2;">
        <p style="margin-bottom: 8px;text-indent: 2.1429em;text-wrap: wrap;"><strong>摘&nbsp;&nbsp;要：</strong>{paper_info["摘要"]}</p>
        <p style="text-indent: 2.1429em;text-wrap: wrap;"><strong>关键词：</strong>{paper_info["关键词"]}</p>
                        </section>
                    </section>
                </section>{paper_info["正文"]}
        <p style="text-wrap: wrap;"><br></p>
                <section style="margin-top: -13px;">
                    <p style="text-wrap: wrap;"><br></p>
                </section>
                <section style="font-size: 14px;color: rgb(131, 131, 131);">
                    <p style="text-wrap: wrap;">转载时烦请注明“转载自北方法学公众号”。</p>
                </section>
                <section style="display: flex;flex-flow: row;margin: 10px 0%;text-align: left;justify-content: flex-start;">
                    <section style="display: inline-block;vertical-align: bottom;width: auto;flex: 100 100 0%;align-self: flex-end;height: auto;">
                        <section style="margin-right: 0%;margin-left: 0%;">
                            <section style="background-color: rgb(101, 66, 54);height: 1px;">
                                <svg viewBox="0 0 1 1"
                                     style="float:left;line-height:0;width:0;vertical-align:top;"></svg>
                            </section>
                        </section>
                    </section>
                    <section style="display: inline-block;vertical-align: bottom;width: auto;align-self: flex-end;flex: 0 0 0%;height: auto;margin-right: 6px;margin-bottom: -3px;margin-left: 6px;">
                        <section style="font-size: 0px;margin-right: 0%;margin-bottom: 1px;margin-left: 0%;transform: translate3d(1px, 0px, 0px);text-align: center;justify-content: center;display: flex;flex-flow: row;">
                            <section style="display: inline-block;width: 22px;vertical-align: top;flex: 0 0 auto;height: auto;background-color: rgb(101, 66, 54);align-self: flex-start;">
                                <section style="margin-right: 0%;margin-bottom: -2px;margin-left: 0%;line-height: 0;">
                                    <section style="vertical-align: middle;display: inline-block;line-height: 0;">
                                        <img data-imgfileid="100001932"
                                             data-ratio="0.74"
                                             data-s="300,640"
                                             data-type="gif"
                                             data-w="300"
                                             style="vertical-align: middle;width: 100%;height: auto !important;"
                                             src="https://mmbiz.qpic.cn/mmbiz_gif/iab7cNlld8FfIY5XjDNqQZJxbCxQLpicUFEyFxYFWNaFibYiaZ747newlxbJbBFO7icxcswX2ibSEzlxWiav9jGqQickTg/640?wx_fmt=gif&amp;from=appmsg">
                                    </section>
                                </section>
                            </section>
                        </section>
                    </section>
                    <section style="display: inline-block;vertical-align: bottom;width: auto;align-self: flex-end;flex: 100 100 0%;">
                        <section style="margin-right: 0%;margin-left: 0%;">
                            <section style="background-color: rgb(101, 66, 54);height: 1px;">
                                <svg viewBox="0 0 1 1"
                                     style="float:left;line-height:0;width:0;vertical-align:top;"></svg>
                            </section>
                        </section>
                    </section>
                </section>
                <section style="margin-bottom: 3px;">
                    <section style="font-size: 12px;">
                        <p style="text-align: center;text-wrap: wrap;">编辑｜谢昕烨</p>
                    </section>
                </section>
                <section style="font-size: 12px;">
                    <p style="text-align: center;text-wrap: wrap;">审核｜赵　亮</p>
                </section>
                <section style="text-align: center;margin-top: 10px;margin-bottom: 10px;line-height: 0;">
                    <section style="vertical-align: middle;display: inline-block;line-height: 0;">
                        <img class="rich_pages wxw-img"
                             data-imgfileid="100001934"
                             data-ratio="0.5555555555555556"
                             data-s="300,640"
                             data-type="png"
                             data-w="1080"
                             style="vertical-align: middle;width: 100%;height: auto !important;"
                             src="https://mmbiz.qpic.cn/mmbiz_png/iab7cNlld8FfIY5XjDNqQZJxbCxQLpicUFmUibj95jJkibqzkySoc9KeWCcdnzy7JymibVtO2FfXvrlTtbRIzDkAFNQ/640?wx_fmt=png&amp;from=appmsg">
                    </section>
                </section>
            </section>
        </body>
        """

        output = """<!DOCTYPE html><html style="padding-left: 95px; padding-right: 95px;" xmlns="http://www.w3.org/1999/xhtml"><head>
            <style type="text/css">
                body{font-family:sans-serif;}
            </style>
            <link rel="stylesheet"
                  type="text/css"
                  href="https://res.wx.qq.com/mpres/zh_CN/htmledition/js/default~media/get_article_structure_tmpl~media/iframe_editor_tmpl.85157959.css">
            <style id="table">
                .selectTdClass{background-color:#edf5fa !important}table.noBorderTable td,table.noBorderTable th,table.noBorderTable caption{border:1px dashed #ddd !important}table{margin-bottom:10px;border-collapse:collapse;display:table;}td,th{padding: 5px 10px;border: 1px solid #DDD;}caption{border:1px dashed #DDD;border-bottom:0;padding:3px;text-align:center;}th{border-top:2px solid #BBB;background:#F7F7F7;}.ue-table-interlace-color-single{ background-color: #fcfcfc; } .ue-table-interlace-color-double{ background-color: #f7faff; }td p{margin:0;padding:0;}
            </style>
            <style id="list">
                ol,ul{margin:0;padding:0;-webkit-box-sizing:border-box;box-sizing:border-box;width:99.9%}li{clear:both;}.list-paddingleft-1 {padding-left:1.2em}
                .list-paddingleft-2 {padding-left:2.2em}
                .list-paddingleft-3 {padding-left:3.2em}
            </style>
        </head>""" + output_body + """</html>"""

        output.replace("__NOT_MOUNT__", "")

        return output

    @staticmethod
    def extract_all_text(docx_file, txt_file):
        # 读取 docx 文件
        doc = Document(docx_file)

        # 打开要写入的 txt 文件，使用 utf-8 编码
        with open(txt_file, 'w', encoding='utf-8') as f:
            # 遍历文档中的段落，将其文字写入 txt 文件
            for para in doc.paragraphs:
                if para.text.strip():  # 检查是否是空白行
                    f.write(para.text + '\n')  # 写入段落内容，并换行

            # 遍历文档中的表格，提取表格中的文字
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():  # 检查是否是空单元格
                            f.write(cell.text + '\n')  # 写入表格单元格内容，并换行

    def parse_file(self, file_path):
        PaperParser.extract_all_text(file_path, "/tmp/d2t_output.txt")
        # 读取文件并解析
        with open("/tmp/d2t_output.txt", 'r', encoding='utf-8') as file:
            paper_text = file.read()
        return self.extract_paper_info(paper_text)

    @staticmethod
    def save_html(html_content, output_path='output.html'):
        # 保存HTML到文件
        with open(output_path, 'w', encoding='utf-16') as file:
            file.write(html_content)


def main():
    parser = argparse.ArgumentParser(description='解析学术论文并生成 HTML 输出。')
    parser.add_argument('input_file', help='输入文本文件的路径')
    parser.add_argument('-o', '--output', default='output.html', help='输出 HTML 文件的路径')
    args = parser.parse_args()

    paper_parser = PaperParser()
    paper_info = paper_parser.parse_file(args.input_file)
    if paper_info:
        html_content = paper_parser.generate_html(paper_info)
        paper_parser.save_html(html_content, args.output)

if __name__ == "__main__":
    main()